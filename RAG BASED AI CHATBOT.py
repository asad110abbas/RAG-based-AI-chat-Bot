from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from docx import Document as DocxDocument
from langchain_openai import ChatOpenAI
import gradio as gr
import os

FAQ_PATH = "FAQs.docx"
DB_PATH = "rag_data"
FINAL_N = 7
FORCE_REBUILD = False

os.environ["OPENAI_API_BASE"] = "https://openrouter.ai/api/v1"
os.environ["OPENAI_API_KEY"] = os.environ.get("OPENAI_API_KEY", "")

docx = DocxDocument(FAQ_PATH)
documents = [Document(page_content=p.text.strip()) for p in docx.paragraphs if p.text.strip()]
if not documents:
    raise ValueError("No content found in FAQs.docx")

embedding = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")

def build_or_load_vectordb():
    need_build = FORCE_REBUILD or (not os.path.exists(DB_PATH)) or (not os.listdir(DB_PATH))
    if need_build:
        vectordb_local = Chroma.from_documents(documents, embedding, persist_directory=DB_PATH)
        vectordb_local.persist()
        return vectordb_local
    else:
        return Chroma(persist_directory=DB_PATH, embedding_function=embedding)

vectordb = build_or_load_vectordb()

llm = ChatOpenAI(model="openai/gpt-5-mini", temperature=0.5, max_tokens=3000)

PROMPT_TEMPLATE = os.environ.get("SYSTEM_PROMPT", "Sorry, I don’t have this info right now. Please visit: https://example.com or contact support.")
prompt = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)

def get_context(query, final_n=FINAL_N):
    dense_results = vectordb.similarity_search_with_score(query, k=final_n)
    if not dense_results:
        return "", False
    top_chunks = [doc.page_content.strip() for doc, _ in dense_results if doc.page_content.strip()]
    return " ".join(top_chunks), True

def run_rag(query):
    context, ok = get_context(query)
    if not ok or not context:
        return "Sorry, I don’t have this info right now. Please visit: https://example.com or contact support."
    final_prompt_messages = prompt.format_messages(context=context, question=query)
    response = llm.invoke(final_prompt_messages)
    return response.content

CSS = """
.gradio-container, body { background-color: #49A8FF !important; font-family: "Segoe UI", Arial, sans-serif; }
#topbar { text-align: center; background: #49A8FF; padding: 12px; border: 2px solid #FCB001; border-radius: 0 0 12px 12px; font-size: 30px; font-weight: bold; color: white; }
#chat-container { background: white; border: 2px solid #FCB001; border-radius: 12px; height: 550px; display: flex; flex-direction: column; justify-content: space-between; overflow: hidden; opacity: 0.95; }
#chatbox { flex: 1; overflow-y: auto; padding: 15px; }
.gr-chatbot .message.user { background: #FFD700 !important; color: black !important; border-radius: 12px 12px 0 12px !important; margin: 5px 0 !important; padding: 10px 14px !important; align-self: flex-end !important; max-width: 80%; }
.gr-chatbot .message.bot { background: #fefefe !important; color: black !important; border: 2px solid #FCB001 !important; border-radius: 12px 12px 12px 0 !important; margin: 5px 0 !important; padding: 10px 14px !important; align-self: flex-start !important; max-width: 80%; }
#input-area { display: flex; align-items: center; gap: 8px; padding: 8px; border-top: 2px solid #FCB001; background: white; }
#msgbox { flex: 1; border: 2px solid #FCB001 !important; border-radius: 8px; padding: 8px; background: white !important; color: black !important; }
#msgbox::placeholder { color: #888; }
#send-btn { background: #fefefe !important; border: 2px solid #FCB001 !important; border-radius: 8px; width: 60px; height: 38px; font-size: 16px; font-weight: bold; cursor: pointer; }
#send-btn:hover { filter: brightness(0.95); }
"""

def chat_fn(message, history):
    history = history + [(message, None)]
    reply = run_rag(message)
    history[-1] = (message, reply)
    return history, ""

with gr.Blocks(css=CSS) as demo:
    gr.HTML("<div id='topbar'>KnowledgeBot AI Chat</div>")
    with gr.Column(elem_id="chat-container"):
        chatbot = gr.Chatbot(elem_id="chatbox", label="")
        with gr.Row(elem_id="input-area"):
            msg = gr.Textbox(placeholder="Type your message...", elem_id="msgbox", show_label=False)
            send = gr.Button("➤", elem_id="send-btn")
    send.click(chat_fn, inputs=[msg, chatbot], outputs=[chatbot, msg])
    msg.submit(chat_fn, inputs=[msg, chatbot], outputs=[chatbot, msg])

demo.launch()
